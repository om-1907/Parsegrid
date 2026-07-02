import os
import re
import logging
import fitz  # PyMuPDF
import docx
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

class DocumentReaderError(Exception):
    """Base exception for document reader module errors."""
    pass

class DocumentCorruptedError(DocumentReaderError):
    """Raised when a document is corrupted, unreadable, or missing."""
    pass

class DocumentEmptyError(DocumentReaderError):
    """Raised when a document contains no extractable text."""
    pass

def clean_text(text: str) -> str:
    """
    Cleans raw extracted text:
    1. Replaces 3 or more consecutive newlines with exactly 2 newlines.
    2. Replaces 2 or more consecutive horizontal spaces with a single space.
    3. Strips leading/trailing whitespace.
    """
    if not text:
        return ""
    cleaned = re.sub(r'\n{3,}', '\n\n', text)
    cleaned = re.sub(r'[ \t]{2,}', ' ', cleaned)
    return cleaned.strip()

def extract_pdf_text(file_path: str) -> str:
    """Extracts text from PDF using PyMuPDF."""
    try:
        doc = fitz.open(file_path)
    except Exception as e:
        raise DocumentCorruptedError(
            f"Failed to open or read the PDF file at '{file_path}'. "
            f"It may be corrupted, missing, or unsupported. Details: {e}"
        ) from e

    extracted_text_parts = []
    try:
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            if text:
                extracted_text_parts.append(text)
    except Exception as e:
        doc.close()
        raise DocumentCorruptedError(
            f"An error occurred while reading pages of the PDF file '{file_path}'. Details: {e}"
        ) from e
    finally:
        doc.close()

    raw_text = "\n".join(extracted_text_parts)
    cleaned = clean_text(raw_text)
    if not cleaned:
        raise DocumentEmptyError(f"The PDF file at '{file_path}' contains no extractable text.")
    return cleaned

def extract_docx_text(file_path: str) -> str:
    """Extracts text from DOCX paragraphs and tables."""
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        raise DocumentCorruptedError(
            f"Failed to open or read the DOCX file at '{file_path}'. Details: {e}"
        ) from e

    full_text = []
    try:
        # Read paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        # Read tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
    except Exception as e:
        raise DocumentCorruptedError(
            f"An error occurred while reading DOCX structure for '{file_path}'. Details: {e}"
        ) from e

    raw_text = "\n\n".join(full_text)
    cleaned = clean_text(raw_text)
    if not cleaned:
        raise DocumentEmptyError(f"The DOCX file at '{file_path}' contains no extractable text.")
    return cleaned

def extract_plain_text(file_path: str) -> str:
    """Extracts text from TXT, MD, or CSV plain text files."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            raw_text = f.read()
    except Exception as e:
        raise DocumentCorruptedError(
            f"Failed to read text file at '{file_path}'. Details: {e}"
        ) from e

    cleaned = clean_text(raw_text)
    if not cleaned:
        raise DocumentEmptyError(f"The text file at '{file_path}' contains no extractable text.")
    return cleaned

def extract_html_text(file_path: str) -> str:
    """Extracts text from HTML file using BeautifulSoup."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            html_content = f.read()
    except Exception as e:
        raise DocumentCorruptedError(
            f"Failed to read HTML file at '{file_path}'. Details: {e}"
        ) from e

    try:
        soup = BeautifulSoup(html_content, "html.parser")
        for element in soup(["script", "style", "header", "footer", "nav"]):
            element.decompose()
        
        # Get clean paragraphs/text elements
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        raw_text = "\n".join(chunk for chunk in chunks if chunk)
    except Exception as e:
        raise DocumentCorruptedError(
            f"Failed to parse HTML structure for '{file_path}'. Details: {e}"
        ) from e

    cleaned = clean_text(raw_text)
    if not cleaned:
        raise DocumentEmptyError(f"The HTML file at '{file_path}' contains no extractable text.")
    return cleaned

def extract_document_text(file_path: str) -> str:
    """
    Detects file extension and extracts text using the appropriate parser.
    Supports .pdf, .docx, .txt, .md, .csv, .html, .htm.
    """
    if not os.path.exists(file_path):
        raise DocumentCorruptedError(f"File not found at '{file_path}'.")

    _, ext = os.path.splitext(file_path.lower())
    if ext == ".pdf":
        return extract_pdf_text(file_path)
    elif ext == ".docx":
        return extract_docx_text(file_path)
    elif ext in (".txt", ".md", ".csv"):
        return extract_plain_text(file_path)
    elif ext in (".html", ".htm"):
        return extract_html_text(file_path)
    else:
        raise DocumentCorruptedError(f"Unsupported file extension: '{ext}'")
